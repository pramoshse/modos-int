import streamlit as st
import os
import asyncio
from playwright.async_api import async_playwright
import datetime
import json
import base64
from io import BytesIO

# Librerías para generación de Excel Técnico
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

import subprocess
import sys

@st.cache_resource
def install_playwright_chromium():
    subprocess.run(
        [sys.executable, "-m", "playwright", "install", "chromium"],
        check=True
    )

install_playwright_chromium()

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Piloto HEC v0", layout="wide", initial_sidebar_state="collapsed")

# --- ESTILOS CSS ---
st.markdown("""
    <style>
    .custom-section-header {
        background-color: #B51E2D;
        color: white;
        padding: 10px;
        border-radius: 5px;
        font-size: 18px;
        font-weight: bold;
        margin-bottom: 15px;
        margin-top: 20px;
    }
    .mode-banner {
        padding: 15px;
        border-radius: 8px;
        color: white;
        font-weight: bold;
        text-align: center;
        margin: 10px 0;
        font-size: 16px;
    }
    .card-danger {
        background-color: #fee2e2;
        border-left: 5px solid #ef4444;
        color: #7f1d1d !important;
        padding: 10px;
        margin: 10px 0;
        border-radius: 4px;
    }
    .card-danger strong {
        color: #991b1b !important;
    }
    .final-banner-premium {
        background-color: #B51E2D;
        color: white;
        padding: 20px;
        border-radius: 10px;
        text-align: center;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    .final-banner-premium .label {
        font-size: 14px;
        text-transform: uppercase;
        opacity: 0.8;
    }
    .final-banner-premium .value {
        font-size: 32px;
        font-weight: 900;
        margin-top: 5px;
    }
    </style>
""", unsafe_allow_html=True)

# --- FUNCIONES AUXILIARES ---
def get_base64_image(file):
    if file is None: return None
    try:
        return base64.b64encode(file.getvalue()).decode()
    except:
        return None

def file_to_base64(path):
    if not os.path.exists(path): return None
    try:
        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except:
        return None


# --- GESTIÓN DE BORRADORES JSON ---
DRAFT_VERSION = 1
DRAFT_EXCLUDED_KEYS = {
    "pdf_v5", "excel_v5", "word_v5",
    "draft_upload", "draft_load_requested", "equip_uploader",
    "draft_json_download", "draft_json_filename"
}


def _draft_json_safe(value):
    """Convierte valores de Streamlit/session_state a estructuras compatibles con JSON."""
    if isinstance(value, datetime.datetime):
        return {"__type__": "datetime", "value": value.isoformat()}
    if isinstance(value, datetime.date):
        return {"__type__": "date", "value": value.isoformat()}
    if isinstance(value, (str, int, float, bool)) or value is None:
        return value
    if isinstance(value, list):
        return [_draft_json_safe(item) for item in value]
    if isinstance(value, tuple):
        return [_draft_json_safe(item) for item in value]
    if isinstance(value, dict):
        return {str(k): _draft_json_safe(v) for k, v in value.items()}
    return None


def _draft_restore_value(value):
    """Reconstruye valores especiales, como fechas, al cargar un borrador JSON."""
    if isinstance(value, dict) and "__type__" in value and "value" in value:
        if value["__type__"] == "date":
            try:
                return datetime.date.fromisoformat(value["value"])
            except Exception:
                return datetime.date.today()
        if value["__type__"] == "datetime":
            try:
                return datetime.datetime.fromisoformat(value["value"])
            except Exception:
                return datetime.datetime.now()
    if isinstance(value, list):
        return [_draft_restore_value(item) for item in value]
    if isinstance(value, dict):
        return {k: _draft_restore_value(v) for k, v in value.items()}
    return value


def _ensure_widget_value(key, options, multi=False):
    """Mantiene coherentes los valores restaurados cuando un selector depende de otro campo."""
    if key not in st.session_state:
        return
    options = list(options)
    if multi:
        current_values = st.session_state.get(key) or []
        st.session_state[key] = [value for value in current_values if value in options]
    elif options and st.session_state.get(key) not in options:
        st.session_state[key] = options[0]


def _capture_equipment_photo(uploaded_file=None):
    """Guarda la foto actual del equipo como base64 para que pueda recuperarse desde JSON."""
    if uploaded_file is not None:
        try:
            st.session_state["equip_photo_b64"] = base64.b64encode(uploaded_file.getvalue()).decode()
            st.session_state["equip_photo_name"] = getattr(uploaded_file, "name", "foto_equipo")
            st.session_state["equip_photo_type"] = getattr(uploaded_file, "type", "image/png")
        except Exception:
            pass


def _get_equipment_photo_payload(uploaded_file=None):
    """Obtiene la foto del equipo, desde el archivo cargado actual o desde una foto restaurada."""
    _capture_equipment_photo(uploaded_file)
    photo_b64 = st.session_state.get("equip_photo_b64")
    if not photo_b64:
        return None
    return {
        "filename": st.session_state.get("equip_photo_name", "foto_equipo"),
        "mime_type": st.session_state.get("equip_photo_type", "image/png"),
        "content_base64": photo_b64,
    }


def _build_draft_payload(evaluaciones_fine=None, energias_seleccionadas=None, lista_peligros_final=None, modo_inicial=None, modo_final=None, uploaded_file=None):
    """Compila un borrador JSON completo con todos los widgets y la matriz de riesgos calculada."""
    widgets = {}
    for key, value in st.session_state.items():
        if key in DRAFT_EXCLUDED_KEYS or key.startswith("_draft_"):
            continue
        safe_value = _draft_json_safe(value)
        if safe_value is not None:
            widgets[key] = safe_value

    return {
        "draft_version": DRAFT_VERSION,
        "created_at": datetime.datetime.now().isoformat(timespec="seconds"),
        "app": "Piloto HEC v0",
        "widgets": widgets,
        "equipment_photo": _get_equipment_photo_payload(uploaded_file),
        "computed": {
            "modo_inicial": modo_inicial,
            "modo_final": modo_final,
            "energias_seleccionadas": _draft_json_safe(energias_seleccionadas or {}),
            "lista_peligros_final": _draft_json_safe(lista_peligros_final or []),
            "evaluaciones_fine": _draft_json_safe(evaluaciones_fine or []),
        }
    }


def _apply_draft_payload(payload):
    """Carga en session_state los valores guardados en un borrador JSON."""
    if not isinstance(payload, dict):
        raise ValueError("El archivo JSON no contiene un objeto válido.")

    widgets = payload.get("widgets", {})
    if not isinstance(widgets, dict):
        raise ValueError("El archivo JSON no contiene una sección 'widgets' válida.")

    for key, value in widgets.items():
        if key in DRAFT_EXCLUDED_KEYS or key.startswith("_draft_") or key == "equip_uploader":
            continue
        st.session_state[key] = _draft_restore_value(value)

    photo = payload.get("equipment_photo") or {}
    if isinstance(photo, dict) and photo.get("content_base64"):
        st.session_state["equip_photo_b64"] = photo.get("content_base64")
        st.session_state["equip_photo_name"] = photo.get("filename", "foto_equipo")
        st.session_state["equip_photo_type"] = photo.get("mime_type", "image/png")


import streamlit as st
from PIL import Image
import io

# --- 1. CONFIGURACIÓN E INICIALIZACIÓN ---
# Asegúrate de tener arca.png en la carpeta del proyecto
logo_path = "arca.png" 

st.title("MODOS DE INTERVENCIÓN Y ANÁLISIS DE RIESGOS")
st.markdown(
    """
    <p style='font-size: 0.82rem; color: #64748B; margin-top: -12px; margin-bottom: 18px;'>
        Prueba Piloto: Esta aplicación corresponde a una prueba conceptual (PoC) alojada en infraestructura cloud externa. La información ingresada será utilizada exclusivamente para el procesamiento temporal y generación de documentos asociados a la sesión.
    </p>
    """,
    unsafe_allow_html=True,
)


# --- SIDEBAR: CARGA DE BORRADOR JSON ---
with st.sidebar:
    st.markdown("### Gestión de Borradores")
    draft_upload = st.file_uploader("Buscar archivo borrador", type=["json"], key="draft_upload")
    if st.button("Cargar datos del borrador", use_container_width=True):
        if draft_upload is None:
            st.warning("Seleccioná un archivo .json antes de cargar el borrador.")
        else:
            try:
                draft_payload = json.loads(draft_upload.getvalue().decode("utf-8"))
                _apply_draft_payload(draft_payload)
                st.success("Borrador cargado correctamente. Los campos se restaurarán automáticamente.")
                st.rerun()
            except Exception as exc:
                st.error(f"No se pudo cargar el borrador JSON: {exc}")


# --- 2. DATOS DESPLEGABLES ---
datos_maquinas = ["Aéreos", "Agrupadora", "Alimentador de cartón", "Alimentador de tapa", "Almacén", "Almacén de paleta", "Analizador", "Aplicador", "Aplicador de película", "Bailer", "Banda", "Batch Blend", "Bifurcador", "Bomba de incendios", "Bombas", "Bombeo", "Brazo", "Buffer", "Caldera", "Calentador", "Canulador", "Capsuladora", "Caracol", "Carboblending", "Carbocooler", "Carbonatador", "Carborproporcionador", "Chiller", "Codificador", "Codificador de lata", "Colocador de cánula", "Colocadora", "Combinador de paquetes", "Compactador", "Compresor", "Condensador", "Controlador de etiquetas", "Controlador de línea", "Coronado", "Cuarto frío", "Decapsulador", "Desempacadora", "Desencajonadora", "Desenfardadora", "Desenroscadora", "Desetiquetadora", "Despaletizadora", "Destapadora", "Disolución continua", "Disolución de carbón activado", "Distribuidor", "Divergente", "Divisor", "Divisora de Paquetes", "Dosificador", "Elevador", "Elevador de Tapas", "Empacadora", "Emplayadora", "Encajonadora", "Encogedora", "Enjuagadora", "Ensobretadora", "Entrada", "Envasadora", "Envolvedora", "Equipo Auxiliar", "Equipo TA", "Esmeril", "Estación", "Etiquetadora", "Extractor", "Filtro", "Flejadora", "Formadora", "Helix", "Hicone Applicator", "Homogenizador", "Horizontal Bailer", "Horno", "Humas", "Impresora", "Inspector", "Inspector de Botellas", "Inspector de Cajas", "Inspector de Etiquetas", "Inspector de Máquina de Lata", "Inspector de Nivel", "Inspector de Vacío", "Inspector Electrónico", "Intercambiador", "Intercambiador de Calor", "Inyector", "Lámpara", "Lantech 1", "Lavadora", "Lavadora de cajas", "Lavadora de garrafón", "Lente", "LGV", "Llenadora", "Llenadora Tetrapack", "Llenadora y Rinser", "Magazine de Cajas", "Magazine de Pallets", "Marmita", "Mesa", "Mesa Acumuladora", "Mesa de Descarga", "Mezclador", "Mixer", "Motor", "Nitrógeno", "Ordenador", "Ordenador de cánulas", "Ósmosis", "Ozono", "Paletizadora", "Pasteurizador", "Polipasto", "Posicionador de Paquetes", "PP D001", "PP D002", "PP D003", "PP D004", "PP D005", "PP D006", "PP D007", "PP D008", "PP D009", "PP D010", "PP D011", "PP D012", "Precintador", "Prensa", "Preparador", "Proporcionador", "Rampa", "Rechazo", "Refrigeración", "Revisador", "Rinser", "Robot", "Rolador", "Roscadora", "Ruta", "Salida", "Secador", "Separador", "Sierra", "Sistema de NH3", "Sistema Ergonómico", "Sistema PTAN", "Soldadora", "Soplador", "Straw", "Tablero Eléctrico", "Taladro", "Tanque", "Tanque de disolución", "Tanque de dosage", "Taponadora", "Termoencogible", "Termoformadora", "Tolva", "Torre Enfriamiento", "Transformador", "Transporte de Pallets", "Transporte de Packs", "Transporte de Cajas", "Transportador", "Transporte de Botellas", "Transporte de Preformas", "Trechas", "Triturador", "Túnel de Enfriamiento", "Vertical Bailer", "Pulmón"]

opciones_tareas = {
    "Operaciones normales (0 acceso): Protecciones instaladas de conformidad con OSH-RQ-185": [
        "Operación normal de equipos", "Monitoreo de parámetros operacionales", "Supervisión de producción", 
        "Verificación de alarmas y tendencias", "Cambio de setpoints autorizados", "Arranque normal de equipos", 
        "Detención normal de equipos", "Operación desde HMI"
    ],
    "Intervención Menor (trabajo a través o dentro de áreas protegida): rutinaria, repetitiva e integral": [
        "Acomodo de cartón dañado", "Acomodo de paquetes fuera de posición", "Acomodo de tarima fuera de posición", "Ajuste de clutch", "Ajuste de espreas", "Ajuste de guías", "Ajuste de gusano de entrada", "Ajuste de pata de gallo", "Ajuste de placas", "Ajuste de presión de aire", "Ajuste de sensores", "Ajuste de tiempos de manejo", "Ajuste de torque", "Ajuste de válvulas", "Ajuste de ventosas", "Ajuste por cambio de presentación", "Alineación de sensores", "Aplicación de vortex", "Cambio de bandas", "Cambio de baleros", "Cambio de botella dañada", "Cambio de componentes mecánicos", "Cambio de copas", "Cambio de filtros", "Cambio de guías", "Cambio de manejos", "Cambio de rollo de termoencogible", "Cambio de sistema de arrastre", "Cambio de válvulas", "Cambio de tubos de ventila", "Colocación de tulipas o botellas falsas", "Corte de producto", "Corte manual de línea", "Destrabe de botellas", "Destrabe de cajas", "Inspección ultrasónica", "Levantamiento de botellas caídas", "Levantamiento de paquetes atorados", "Limpieza de áreas de transporte", "Limpieza de cámaras interiores", "Limpieza de componentes mecánicos", "Limpieza de cortinas", "Limpieza de filtros y conexiones", "Limpieza de mesa de corte", "Limpieza de mesa de transferencia", "Limpieza de sensores", "Limpieza de sopladores", "Limpieza de ventiladores", "Limpieza interior de equipos", "Lubricación centralizada", "Lubricación de bandas", "Lubricación de elevadores", "Lubricación de rodillos", "Lubricación de sistemas mecánicos", "Lubricación general", "Operación manual de válvulas", "Predictivo por ultrasonido", "Procedimiento por explosión de botella", "Reemplazo de componentes desgastados", "Reposición de componentes de transportador", "Retiro de acumulación de cartón", "Retiro de acumulación de tapas", "Retiro de botellas defectuosas", "Retiro de botellas sin envolver", "Retiro de etiquetas acumuladas", "Retiro de etiquetas mal aplicadas", "Retiro de plástico y residuos", "Retiro de producto defectuoso", "Retiro de tarima dañada", "Revisión de bandas", "Revisión de cangilones", "Revisión de componentes mecánicos", "Revisión de conexiones", "Revisión de encoder", "Revisión de estrellas de sujeción", "Revisión de filtros", "Revisión de guardas de seguridad", "Revisión de hornilla", "Revisión de intercambiadores", "Revisión de reductores", "Revisión de rodillos", "Revisión de sensores", "Revisión de separadores", "Revisión de sistemas de arrastre", "Revisión de sistemas de seguridad", "Revisión de sopladores", "Revisión de tensores", "Revisión de transportadores", "Revisión de túnel de encogimiento", "Revisión de unidades de mantenimiento", "Revisión de ventiladores", "Verificación de bandas deslizantes", "Verificación de componentes eléctricos", "Verificación de conexiones", "Verificación de guardas protectoras", "Verificación de inocuidad", "Verificación de micro switches", "Verificación de presión", "Verificación de seguridad operacional", "Verificación de sistemas neumáticos", "Verificación de transportadores"
    ]
}

# --- 3. INTERFAZ ---
col1, col2 = st.columns(2)

with col1:
    st.markdown('<div class="custom-section-header">Información de Sitio</div>', unsafe_allow_html=True)
    lista_sitios = [
        "Planta Tucumán", "Planta Formosa", "Planta Salta",
        "CEDI Corrientes", "CEDI Goya", "CEDI Posadas", "CEDI Resistencia", "CEDI Saenz Peña",
        "Ingenio Famaillá", "Ingenio Bella Vista", "CEDI J.V.Gonzalez", "CEDI Ledesma",
        "CEDI Metán", "CEDI Orán", "CEDI Jujuy", "CEDI San Pedro", "CEDI Sgo Del Estero",
        "CEDI Tartagal", "CEDI Chilecito", "CEDI Concepción", "CEDI La Rioja", "CEDI Catamarca"
    ]
    sitio = st.selectbox("Sitio", lista_sitios, key="sitio")
    
    # Clasificación de Negocio
    negocio = "Bebidas"
    tipo_sitio = "Planta" if "Planta" in sitio else "CEDI"
    if "Ingenio" in sitio:
        negocio = "Ingenios"
        tipo_sitio = "Ingenio"
        
    area_sector = st.selectbox("Área / Sector", [
        "Producción", "Mantenimiento", "Calidad", "Servicios Auxiliares", 
        "Expedición", "Comercial", "Logística", "Automotor"
    ], key="area_sector")
    
    linea = "N/A"
    maquina = "N/A"
    subsector_equipo = "N/A"
    
    if area_sector == "Producción":
        # Lógica de líneas según planta
        if sitio == "Planta Salta":
            linea_options = ["Línea 01", "Línea 02", "Línea 03", "Línea 04", "Línea 05", "Línea 06", "Línea Dual Pack", "Línea 1000 Tetra", "Línea 200 Tetra", "Línea Bidones", "Sorting", "BBOX"]
        elif sitio == "Planta Tucumán":
            linea_options = ["Línea 01", "Línea 02", "Línea 03", "Línea 04", "Línea 05", "Línea 06", "Sorting", "BBOX"]
        else:
            linea_options = ["Línea 01", "Línea 02", "Línea 03", "Línea 04"]
        _ensure_widget_value("linea", linea_options)
        linea = st.selectbox("Línea de Embotellado", linea_options, key="linea")
        maquina = st.selectbox("Máquina", datos_maquinas, key="maquina")
    else:
        subsector_equipo = st.text_input("Subsector / Equipo o Sistema", key="subsector_equipo")
    
    st.markdown('<div class="custom-section-header">Datos de Equipo</div>', unsafe_allow_html=True)
    fabricante = st.text_input("Fabricante", key="fabricante")
    modelo = st.text_input("Modelo / Marca", key="modelo")
    anio = st.number_input("Año de fabricación", min_value=1950, max_value=2030, step=1, key="anio")

with col1:
    st.markdown('<div class="custom-section-header">Clasificación de Tarea</div>', unsafe_allow_html=True)
    clasificacion = st.radio("Clasifique tarea(s):", list(opciones_tareas.keys()) + ["Mantenimiento", "Equipo energizado (modo jog o enseñanza del equipo)"], key="clasificacion")

    tareas_predefinidas = []
    if clasificacion in opciones_tareas:
        _ensure_widget_value("tareas_predefinidas", opciones_tareas[clasificacion], multi=True)
        tareas_predefinidas = st.multiselect("Seleccione tarea(s):", opciones_tareas[clasificacion], key="tareas_predefinidas")
    
    tareas_manuales = st.text_area("Identifique tareas adicionales (Enter para nuevas líneas):", key="tareas_manuales")
    
    # Combinar tareas para el procesamiento posterior
    lista_tareas_combinadas = []
    if tareas_predefinidas:
        lista_tareas_combinadas.extend(tareas_predefinidas)
    if tareas_manuales:
        lista_tareas_combinadas.extend([t.strip() for t in tareas_manuales.split("\n") if t.strip()])
    
    frecuencia = st.selectbox("Frecuencia promedio", [
        "De 15 a 20 veces por hora", "De 5 a 20 veces por turno", "De 1 a 15 veces por turno",
        "De 3 a 10 veces por turno", "De 3 a 6 veces por turno", "De 2 a 4 veces por turno",
        "De 1 a 5 veces por turno", "De 1 a 3 veces por turno", "De 1 a 2 veces por turno",
        "De 1 a 2 veces por día", "De 3 a 4 veces por semana", "De 1 a 3 veces por semana",
        "1 vez por semana", "De 1 a 3 veces por mes", "De 1 a 2 veces por mes",
        "1 vez cada 2 meses", "1 vez cada 3 meses", "1 vez cada 4 meses",
        "1 vez cada 6 meses", "1 vez al año", "1 vez cada 2 años",
        "1 vez cada 3 años", "1 vez cada 4 años", "1 vez cada 5 años"
    ], key="frecuencia")
    duracion = st.selectbox("Duración promedio", [
        "1 minuto", "2 minutos", "3 minutos", "5 minutos", "10 minutos", "15 minutos",
        "25 minutos", "30 minutos", "1 hora", "2 horas", "3 horas", "4 horas",
        "5 horas", "7 horas", "8 horas", "De 9 a 12 horas", "Más de 12 horas"
    ], key="duracion")
    
    tareas = lista_tareas_combinadas

with col2:
    st.markdown('<div class="custom-section-header">Carga de Foto</div>', unsafe_allow_html=True)
    equip_file = st.file_uploader("Subir Foto del Equipo", type=["png", "jpg", "jpeg"], key="equip_uploader")
    if equip_file:
        _capture_equipment_photo(equip_file)
        st.image(equip_file, caption="Equipo Seleccionado", use_container_width=True)
    elif st.session_state.get("equip_photo_b64"):
        try:
            st.image(BytesIO(base64.b64decode(st.session_state["equip_photo_b64"])), caption="Equipo Seleccionado (borrador)", use_container_width=True)
        except Exception:
            st.warning("La foto guardada en el borrador no pudo visualizarse.")

    st.markdown('<div class="custom-section-header">Identificación de Modo Inicial de Intervención</div>', unsafe_allow_html=True)
    nivel_acceso = st.radio(
        "Nivel de Acceso / Exposición Requerido",
        [
            "La tarea requiere que el equipo esté energizado mientras se trabaja en o alrededor de zonas peligrosas de equipo",
            "La tarea incluye desmontaje/mantenimiento o contacto con partes energizadas",
            "La tarea requiere ACCESO CORPORAL COMPLETO (trabajar DENTRO DE GUARDA INTERBLOQUEADO)",
            "La tarea requiere ACCESO CORPORAL PARCIAL (trabajar A TRAVÉS DE GUARDA INTERBLOQUEADO)",
            "La tarea no requiere NINGÚN TIPO DE ACCESO CORPORAL"
        ],
        key="nivel_acceso"
    )
    
    # --- LÓGICA MODO INICIAL ---
    modo_inicial = "Modo 0"
    if "equipo esté energizado" in nivel_acceso:
        modo_inicial = "Modo 4"
    elif "desmontaje/mantenimiento" in nivel_acceso:
        modo_inicial = "Modo 3"
    elif "ACCESO CORPORAL COMPLETO" in nivel_acceso:
        modo_inicial = "Modo 2"
    elif "ACCESO CORPORAL PARCIAL" in nivel_acceso:
        modo_inicial = "Modo 1"
        
    colores_modos = {
        "Modo 0": "#69C97F",
        "Modo 1": "#00B0F0",
        "Modo 2": "#00B0F0",
        "Modo 3": "#FF560E",
        "Modo 4": "#F30009"
    }
    color_ini = colores_modos.get(modo_inicial, "#B51E2D")
    st.markdown(f'<div class="mode-banner" style="background-color: {color_ini}; color: white;">Modo Inicial Identificado: {modo_inicial}</div>', unsafe_allow_html=True)
    
    # Mensajes de seguridad dinámicos
    mensajes_seguridad = {
        "Modo 0": "Asegúrese de que el ACCESO CERO, la protección de la máquina debe estar en su lugar y en buenas condiciones.",
        "Modo 1": "DETENGA LA MÁQUINA y trabaje a través de la GUARDA CON INTERLOCK.",
        "Modo 2": "DETENGA LA MÁQUINA y aplique 'BLOQUEO' para EVITAR el REINICIO.",
        "Modo 3": "DETENGA LA MÁQUINA y aplique LOTO COMPLETO para realizar la tarea.",
        "Modo 4": "MÁXIMA PRECAUCIÓN: Equipo energizado con acceso a zonas peligrosas."
    }
    
    if modo_inicial in ["Modo 1", "Modo 2"]:
        msg_inicial = "*Modo sujeto a validación en 5. VALIDACIÓN DEL MODO INICIAL"
    else:
        msg_inicial = mensajes_seguridad.get(modo_inicial, "")
        
    if msg_inicial:
        st.markdown(f'<div style="color: {color_ini}; font-weight: bold; border: 2px solid {color_ini}; padding: 10px; border-radius: 5px;">Acción Requerida: {msg_inicial}</div>', unsafe_allow_html=True)
        
    st.markdown('<div style="margin-top: 25px;"></div>', unsafe_allow_html=True)
    evaluador = st.text_input("Nombre y Apellido de Evaluador", key="evaluador")
    puesto_evaluador = st.text_input("Puesto", key="puesto_evaluador")
    fecha_evaluacion = st.date_input("Fecha de Evaluación", value=datetime.date.today(), key="fecha_evaluacion")

# --- FUENTES DE ENERGÍA ---
st.markdown('<div class="custom-section-header">3. Fuentes de Energía(s) para la(s) Tarea(s) / Modo Seleccionado</div>', unsafe_allow_html=True)
fuentes_energia = [
    ("Eléctrica", "380V"), ("Mecánica", "N/A"), ("Hidráulica", "N/A"),
    ("Neumática", "6 bar"), ("Térmica", "N/A"), ("Química", "N/A"),
    ("Potencial", "N/A"), ("Magnética", "N/A"), ("Radiante", "N/A"),
    ("Presión de Agua", "N/A")
]

energias_seleccionadas = {}
cols_eng = st.columns(3)
for i, (nombre, val_def) in enumerate(fuentes_energia):
    with cols_eng[i % 3]:
        activado = st.checkbox(nombre, value=(i==0 or i==3), key=f"chk_{i}")
        if activado:
            if nombre == "Química":
                sustancias_base = ["Amoníaco (NH3)", "Gas carbónico (CO2)", "Soda cáustica (NaOH)", "Nitrógeno", "Gas Oil", "Gas Natural", "GLP"]
                seleccion_sustancias = st.multiselect("Identifique la sustancia(s) química(s):", options=sustancias_base + ["Otra (Manual)"], key="sel_quimica")
                
                detalles_quimica = []
                for s in seleccion_sustancias:
                    if s == "Otra (Manual)":
                        otra_s = st.text_input("Especifique otra sustancia:", key="otra_quimica")
                        mag_s = st.text_input(f"Magnitud para {otra_s}:", key="mag_otra_quimica")
                        if otra_s: detalles_quimica.append(f"{otra_s} ({mag_s})")
                    else:
                        mag_s = st.text_input(f"Magnitud para {s}:", key=f"mag_{s}")
                        detalles_quimica.append(f"{s} ({mag_s})")
                
                if detalles_quimica:
                    energias_seleccionadas[nombre] = " | ".join(detalles_quimica)
            else:
                magnitud = st.text_input(f"Magnitud {nombre}", value=val_def, key=f"mag_{i}")
                energias_seleccionadas[nombre] = magnitud

controles_por_categoria = {
    "Ingeniería": [
        "Interlocks / enclavamientos de puertas", "Sensores ópticos / cortinas de luz / barreras infrarrojas",
        "Paradas de emergencia", "Vallas fijas 360° / cerramientos / mallas metálicas",
        "Protecciones físicas de transmisiones y engranajes", "Plataformas, pasarelas y escaleras antideslizantes",
        "Barandas, rodapiés y travesaños intermedios", "Sistemas de presión positiva",
        "Ventilación industrial / extractores / ventiladores", "Sistemas automáticos de parada",
        "Sistemas automatizados de sanitación o inyección química", "Protección contra proyección de vidrios",
        "Puestas a tierra / diferenciales eléctricos / tensión segura 24V", "Iluminación LED y mejoras de iluminación",
        "Red fija contra incendios", "Alarmas sonoras y balizas luminosas", "Detectores de oxígeno / gases",
        "Separación física peatón–autoelevador", "Diseño ergonómico de puestos",
        "Carros ergonómicos y ayudas mecánicas", "Pisos y superficies antideslizantes",
        "Bandejas de contención y drenajes", "Limitadores de velocidad en autoelevadores"
    ],
    "Administrativas / Procedimientos": [
        "Procedimiento de Bloqueo Operacional", "Procedimiento LOTO / Control de Energías Peligrosas",
        "Procedimiento Protección de Máquinas", "Procedimientos seguros de limpieza, destrabe y sanitación",
        "Procedimientos para trabajos en altura", "Procedimientos eléctricos seguros",
        "Procedimientos para manejo de sustancias peligrosas", "Procedimientos de manejo de contratistas",
        "Permisos de trabajo", "Reglas que Salvan Vidas (RSV)", "Call To Action / SIF / HOP",
        "Programa de Seguridad Basado en el Comportamiento (PSBC)", "Charlas de 5 minutos",
        "Tarjetas de observación", "Reportes e investigación de incidentes",
        "Controles operacionales y rondas de seguridad", "Señalización de advertencia, prohibición y obligación",
        "Demarcación de sendas peatonales", "Reglas para peatones y autoelevadores",
        "Campañas preventivas (golpe de calor, alcoholismo, etc.)", "Simulacros y brigadas de emergencia",
        "Capacitación técnica y de seguridad", "Certificación de operadores y equipos",
        "Planes de mantenimiento preventivo y correctivo", "Estudios ambientales y mediciones"
    ],
    "EPP": [
        "Calzado de seguridad", "Anteojos de seguridad", "Protector facial", "Guantes anticorte",
        "Guantes dieléctricos", "Guantes de PVC", "Mamelucos impermeables", "Protección respiratoria",
        "Protectores auditivos", "Arnés anticaídas", "Chalecos reflectivos", "Casco de seguridad", "Mangas anticorte"
    ],
    "Salud Ocupacional": [
        "Estudios médicos preventivos y periódicos", "Exámenes preocupacionales", "Estudios psicológicos y neurológicos",
        "Vigilancia de salud ocupacional", "Pausas activas / descansos reglamentados", "Control de carga térmica",
        "Control audiométrico", "Seguimiento médico de exposición", "Programas de bienestar y salud mental",
        "Prevención de alcoholismo y drogas"
    ],
    "Gestión de Emergencias": [
        "Planes de emergencia", "Simulacros de evacuación", "Formación de brigadas", "Respuesta a emergencias químicas",
        "Procedimientos ante incendios", "Sistemas de evacuación y señalización de salidas"
    ],
    "Seguridad Patrimonial y Tránsito": [
        "Cámaras de monitoreo", "Guardia privada", "Control de circulación vehicular",
        "Alarmas de retroceso en autoelevadores", "Balizas reflectivas", "Señalización vial interna",
        "Sendas peatonales segregadas", "Control de velocidad vehicular"
    ],
    "Gestión Ergonómica": [
        "Rotación de tareas", "Banquetas ergonómicas", "Regulación de altura de puestos",
        "Ayudas mecánicas para manipulación", "Estudios ergonómicos", "Rediseño de puestos de trabajo",
        "Disminución de manipulación manual de cargas"
    ]
}

peligros_por_categoria = {
    "Mecánicos": [
        "Atrapamiento en partes móviles", "Enredos con elementos rotativos", "Aplastamiento", "Cizallamiento",
        "Golpes contra objetos fijos o móviles", "Proyección de partículas o fragmentos", "Caída de objetos",
        "Contacto con superficies cortantes o punzantes", "Energía mecánica almacenada", "Movimiento inesperado de equipos",
        "Fallas o liberación de presión mecánica"
    ],
    "Eléctricos": [
        "Contacto eléctrico directo", "Contacto eléctrico indirecto", "Arco eléctrico", "Energía eléctrica almacenada",
        "Baja tensión (220V / 380V)", "Equipos energizados"
    ],
    "Físicos": [
        "Ruido elevado (>85 dBA)", "Temperaturas extremas (calor / frío)", "Iluminación insuficiente",
        "Vibraciones", "Radiación no ionizante", "Superficies calientes o frías"
    ],
    "Otros": [
        "Resbalones y tropiezos", "Caídas al mismo nivel", "Caídas a distinto nivel", "Superficies irregulares",
        "Orden y limpieza deficientes", "Deficiencias en accesos o circulación", "Incendio", "Explosión"
    ],
    "Ergonómicos": [
        "Manipulación manual de cargas", "Sobreesfuerzo físico", "Posturas forzadas", "Movimientos repetitivos",
        "Trabajos prolongados de pie", "Diseño inadecuado del puesto de trabajo"
    ],
    "Químicos": [
        "Inhalación de sustancias químicas", "Contacto químico con piel u ojos", "Exposición a vapores, gases o nieblas",
        "Sustancias corrosivas", "Sustancias inflamables", "Derrames químicos"
    ],
    "Transporte / Tránsito": [
        "Choques o colisiones vehiculares", "Interacción peatón–autoelevador", "Estado deficiente de caminos o rutas",
        "Atropellamiento", "Maniobras inseguras"
    ],
    "Psicosociales": [
        "Carga excesiva de trabajo", "Fatiga laboral", "Estrés laboral", "Violencia o acoso laboral",
        "Robos o asaltos", "Consumo de alcohol o sustancias", "Condiciones personales de salud que afecten la tarea"
    ],
    "Especiales": [
        "Espacios confinados", "Trabajo en altura", "Trabajos en caliente", "Atmósferas peligrosas",
        "Deficiencia de oxígeno", "Riesgo biológico", "Condiciones climáticas severas"
    ]
}

dict_prob = {
    0.1: "0.1 - Casi imposible: secuencia o consecuencia prácticamente imposible",
    0.5: "0.5 - Coincidencia extremadamente remota pero concebible",
    1.0: "1.0 - Sería una coincidencia remotamente posible. Se sabe que ha ocurrido. (1%)",
    3.0: "3.0 - Sería una secuencia o coincidencia rara: 10% (Último año)",
    6.0: "6.0 - Es posible: nada extraño tiene una probabilidad del 50% (Últimos 6 meses)",
    10.0: "10.0 - Convicción: es el resultado más probable y esperado"
}
dict_exp = {
    0.5: "0.5 - Muy raramente el riesgo existe, sin antecedente de personal expuesto",
    1.0: "1.0 - Raramente: el riesgo existe y se tiene antecedentes de personal expuesto",
    2.0: "2.0 - Inusual: hasta una vez al año",
    3.0: "3.0 - Ocasionalmente: una vez a la semana o una vez al mes",
    6.0: "6.0 - Frecuentemente: Aproximadamente una o dos veces al día",
    10.0: "10.0 - Continuamente: La situación ocurre continuamente o muchas veces al día"
}
dict_cons = {
    1.0: "1.0 - Heridas leves, contusiones, golpes",
    5.0: "5.0 - Lesiones incapacitantes",
    15.0: "15.0 - Lesiones muy graves, amputaciones, invalidez permanente",
    25.0: "25.0 - Una muerte",
    50.0: "50.0 - Varias muertes o daños materiales graves",
    100.0: "100.0 - Catastróficas: numerosas muertes, grandes daños"
}

def clasificar_fine(gp):
    if gp > 401: return "Riesgo Inminente", "#EF4444"
    elif 201 <= gp <= 400: return "Riesgo Alto", "#F97316"
    elif 71 <= gp <= 200: return "Riesgo Notable", "#EAB308"
    elif 21 <= gp <= 70: return "Riesgo Moderado", "#3B82F6"
    else: return "Riesgo Aceptable", "#10B981"

# --- 4. EVALUACIÓN COMPLETA DE RIESGOS ---
st.markdown('<div class="custom-section-header">4. Identificación de Peligros y Evaluación de Riesgos (W. Fine)</div>', unsafe_allow_html=True)

categorias_seleccionadas = st.multiselect("Seleccione Categoría(s) de Peligro:", options=list(peligros_por_categoria.keys()), key="categorias_seleccionadas")

lista_peligros_final = []
for cat in categorias_seleccionadas:
    pels = st.multiselect(f"Peligros en categoría {cat}:", options=peligros_por_categoria[cat], key=f"sel_pel_{cat}")
    lista_peligros_final.extend(pels)
    p_manual = st.text_input(f"Peligros adicionales manuales para {cat}:", key=f"p_man_{cat}")
    if p_manual:
        lista_peligros_final.append(f"{cat}: {p_manual}")

evaluaciones_fine = []

if lista_peligros_final:
    for i, pel in enumerate(lista_peligros_final):
        st.markdown(f'<div class="card-danger"><strong>⚠️ Peligro #{i+1}:</strong> {pel}</div>', unsafe_allow_html=True)
        
        c_i1, c_i2, c_i3 = st.columns(3)
        with c_i1: p_i = st.selectbox(f"Probabilidad (Inherente) #{i+1}", options=list(dict_prob.keys()), format_func=lambda x: dict_prob[x], key=f"p_i_{i}", index=3)
        with c_i2: e_i = st.selectbox(f"Exposición (Inherente) #{i+1}", options=list(dict_exp.keys()), format_func=lambda x: dict_exp[x], key=f"e_i_{i}", index=4)
        with c_i3: c_i = st.selectbox(f"Consecuencia (Inherente) #{i+1}", options=list(dict_cons.keys()), format_func=lambda x: dict_cons[x], key=f"c_i_{i}", index=2)
        
        gp_inherente = round(p_i * e_i * c_i, 1)
        clase_i, color_i = clasificar_fine(gp_inherente)
        st.markdown(f'<p style="color:{color_i}; font-size:12px; margin-top:-10px;"><strong>GP Inherente: {gp_inherente} ({clase_i}) [P:{p_i} * E:{e_i} * C:{c_i}]</strong></p>', unsafe_allow_html=True)
        
        st.markdown(f"**Asignar Medidas de Control para Peligro #{i+1}:**")
        cat_ctrl_sel = st.multiselect(f"Seleccione Categoría(s) de Control #{i+1}:", options=list(controles_por_categoria.keys()), key=f"cat_ctrl_{i}")
        
        lista_ctrl_final = []
        for c_cat in cat_ctrl_sel:
            c_sel = st.multiselect(f"Medidas en {c_cat} #{i+1}:", options=controles_por_categoria[c_cat], key=f"sel_ctrl_{c_cat}_{i}")
            lista_ctrl_final.extend(c_sel)
            cs_manual = st.text_input(f"Medidas adicionales manuales para {c_cat} (Peligro #{i+1}):", key=f"cs_man_{c_cat}_{i}")
            if cs_manual:
                lista_ctrl_final.append(cs_manual)
        
        controles_seleccionados = lista_ctrl_final
        
        c_r1, c_r2, c_r3 = st.columns(3)
        with c_r1: p_r = st.selectbox(f"Probabilidad (Residual) #{i+1}", options=list(dict_prob.keys()), format_func=lambda x: dict_prob[x], key=f"p_r_{i}", index=1)
        with c_r2: e_r = st.selectbox(f"Exposición (Residual) #{i+1}", options=list(dict_exp.keys()), format_func=lambda x: dict_exp[x], key=f"e_r_{i}", index=2)
        with c_r3: c_r = st.selectbox(f"Consecuencia (Residual) #{i+1}", options=list(dict_cons.keys()), format_func=lambda x: dict_cons[x], key=f"c_r_{i}", index=1)
        
        gp_residual = round(p_r * e_r * c_r, 1)
        clase_r, color_r = clasificar_fine(gp_residual)
        st.markdown(f'<p style="color:{color_r}; font-size:12px; margin-top:-10px;"><strong>GP Residual Mitigado: {gp_residual} ({clase_r}) [P:{p_r} * E:{e_r} * C:{c_r}]</strong></p>', unsafe_allow_html=True)
        st.markdown('---')
        
        evaluaciones_fine.append({
            "peligro": pel, "controles": ", ".join(controles_seleccionados),
            "p_i": p_i, "e_i": e_i, "c_i": c_i, "gp_i": gp_inherente, "clase_i": clase_i,
            "p_r": p_r, "e_r": e_r, "c_r": c_r, "gp_r": gp_residual, "clase_r": clase_r, "color_r": color_r
        })

# --- 5. VALIDACIÓN DEL MODO SEGURO (Solo para Modo 1 y 2) ---
modo_final = modo_inicial
riesgo_no_aceptable = any(item["gp_r"] > 20 for item in evaluaciones_fine)

if modo_inicial in ["Modo 1", "Modo 2"]:
    st.markdown('<div class="custom-section-header">5. Validación del Modo Inicial</div>', unsafe_allow_html=True)
    st.warning(f"Validación requerida para {modo_inicial}")
    
    if modo_inicial == "Modo 1":
        pregunta = "¿Se garantiza control fiable para prevenir rearranque? Enclavamiento (ISO PLd/PLe, ANSI Cat. 3/4)"
        control_fiable = st.radio(pregunta, ["Sí", "No"], index=0, key="control_fiable")
        if riesgo_no_aceptable or control_fiable == "No":
            modo_final = "Modo 3"
            st.error("⚠️ Validación Fallida: Se requiere MODO FINAL 3 debido a riesgo residual > 20 o falta de control fiable.")
            
    elif modo_inicial == "Modo 2":
        pregunta = "¿Se garantiza control fiable para prevenir rearranque? Enclavamiento (ISO PLd/PLe, ANSI Cat. 3/4) + bloqueo físico de puerta o sistema de Llave Atrapada (Trapped o Fortress Key)"
        control_fiable = st.radio(pregunta, ["Sí", "No"], index=0, key="control_fiable")
        if riesgo_no_aceptable or control_fiable == "No":
            modo_final = "Modo 3"
            st.error("⚠️ Validación Fallida: Se requiere MODO FINAL 3 debido a riesgo residual > 20 o falta de control fiable.")

# --- 6. CONCLUSIÓN ---
num_concl = "5" if modo_inicial in ["Modo 0", "Modo 3", "Modo 4"] else "6"
st.markdown(f'<div class="custom-section-header">{num_concl}. Conclusiones / Observaciones</div>', unsafe_allow_html=True)
col_aud1, col_aud2 = st.columns(2)

with col_aud1:
    medidas_conclusiones = st.text_area("Observaciones", value="", key="medidas_conclusiones")

with col_aud2:
    color_fin = colores_modos.get(modo_final, "#B51E2D")
    st.markdown(f"""
        <div class="mode-banner" style="background-color: {color_fin}; font-size: 20px; padding: 20px; color: white;">
            <div style="font-size: 12px; text-transform: uppercase; opacity: 0.8;">Modo Final de Intervención</div>
            {modo_final.upper()}</div>
    """, unsafe_allow_html=True)
    
    msg_final = mensajes_seguridad.get(modo_final, "")
    if msg_final:
        st.markdown(f'<div style="color: {color_fin}; font-weight: bold; border: 2px solid {color_fin}; padding: 15px; border-radius: 5px; text-align: center;">Instrucción Final: {msg_final}</div>', unsafe_allow_html=True)


# --- SIDEBAR: DESCARGA DE BORRADOR JSON ---
with st.sidebar:
    st.markdown("---")
    draft_payload = _build_draft_payload(
        evaluaciones_fine=evaluaciones_fine,
        energias_seleccionadas=energias_seleccionadas,
        lista_peligros_final=lista_peligros_final,
        modo_inicial=modo_inicial,
        modo_final=modo_final,
        uploaded_file=equip_file,
    )
    draft_json = json.dumps(draft_payload, ensure_ascii=False, indent=2)
    st.download_button(
        "Descargar archivo en borrador",
        data=draft_json,
        file_name=f"Borrador_Modos_HEC_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
        mime="application/json",
        use_container_width=True,
    )
    st.caption("El borrador conserva los campos del formulario, la foto del equipo y la matriz de riesgos.")

# --- PLANTILLA HTML AVANZADA ---
HTML_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        @page { size: A4 portrait; margin: 10mm 8mm; }
        body { font-family: 'Bahnschrift', Arial, sans-serif; color: #0f172a; margin: 0; line-height: 1.2; font-size: 8.5pt; }
        .header-table { width: 100%; border-collapse: collapse; margin-bottom: 12px; border: 1.5pt solid #000; }
        .header-table td { border: 0.5pt solid #000; vertical-align: middle; padding: 5px; }
        .header-logo { width: 20%; text-align: center; background-color: #fff; }
        .header-title { width: 55%; text-align: center; background-color: #000; color: #fff; }
        .header-title h1 { font-size: 12pt; margin: 2px 0; text-transform: uppercase; }
        .header-title h2 { font-size: 10pt; margin: 2px 0; font-weight: normal; color: #ccc; }
        .header-title h3 { font-size: 11pt; margin: 2px 0; color: #fff; background-color: #444; padding: 2px; }
        .header-info { width: 25%; font-size: 8pt; background-color: #eee; }
        .header-info table { width: 100%; border-collapse: collapse; }
        .header-info td { border: none; padding: 2px 5px; }
        
        .section { margin-bottom: 10px; border: 1px solid #B51E2D; border-radius: 4px; overflow: hidden; page-break-inside: avoid; }
        .section-header { background-color: #B51E2D; color: #ffffff; padding: 6px 10px; font-size: 9pt; font-weight: bold; text-transform: uppercase; }
        
        .grid-layout { width: 100%; border-collapse: collapse; }
        .grid-layout td { padding: 6px 10px; border: 1px solid #e2e8f0; vertical-align: top; }
        
        .matrix-table { width: 100%; border-collapse: collapse; font-size: 8pt; margin-top: 5px; }
        .matrix-table th { background: #444444; color: white; padding: 5px; text-align: center; font-weight: bold; border: 1px solid #444444; }
        .matrix-table td { padding: 6px; border: 1px solid #cbd5e1; text-align: center; }
        .left-align { text-align: left !important; }
        
        .img-container { text-align: center; padding: 10px; }
        .img-container img { max-width: 100%; height: auto; max-height: 160px; border: 1px solid #cbd5e1; border-radius: 4px; }
        
	        .banner-final-pdf { padding: 10px; text-align: center; border: 2px solid #B51E2D; border-radius: 4px; margin-top: 12px; page-break-inside: avoid; color: white; }
	    </style>
	</head>
	<body>
    <table class="header-table">
        <tr>
            <td class="header-logo">__LOGO_HTML__</td>
            <td class="header-title">
                <h1>CONTROL DE ENERGÍAS PELIGROSAS</h1>
                <h2>__NEGOCIO__ - __PLANTA__</h2>
                <h3>MODOS DE INTERVENCIÓN Y EVALUACIÓN DE RIESGOS</h3>
            </td>
            <td class="header-info">
                <table>
                    <tr><td><strong>CÓDIGO:</strong></td><td></td></tr>
                    <tr><td><strong>REVISIÓN:</strong></td><td></td></tr>
                    <tr><td><strong>FECHA:</strong></td><td>__FECHA_ACTUAL__</td></tr>
                </table>
            </td>
        </tr>
    </table>

    <div class="section">
        <div class="section-header">1. IDENTIFICACIÓN DE EQUIPO Y TAREAS</div>
        <table class="grid-layout">
            <tr>
                <td style="width: 50%;">
                    <strong>Negocio:</strong> __NEGOCIO__ | <strong>Sitio:</strong> __PLANTA__<br>
                    <strong>Ubicación:</strong> __AREA__ — __LINEA__<br>
                    <strong>Equipo/Máquina:</strong> __EQUIPO_DESC__<br>
                    <strong>Fabricante:</strong> __FABRICANTE__ | <strong>Modelo:</strong> __MODELO__ | <strong>Año:</strong> __ANIO__<br>
                    <strong>Clasificación de tarea:</strong> __CLASIFICACION__<br>
                    <strong>Frecuencia:</strong> __FRECUENCIA__ | <strong>Duración:</strong> __DURACION__<br>
                    <strong>Tarea(s):</strong><br>__TAREAS__
                </td>
                <td style="width: 50%; text-align: center;" class="img-container">
                    <strong>REGISTRO VISUAL DEL EQUIPO:</strong><br>
                    __FOTO_EQUIPO__
                </td>
            </tr>
        </table>
    </div>

    <div class="section">
        <div class="section-header">2. IDENTIFICACIÓN DE MODO INICIAL DE INTERVENCIÓN</div>
        <table class="grid-layout">
            <tr>
                <td>
                    <strong>Nivel de Acceso / Exposición Requerido:</strong> __NIVEL_ACCESO__<br>
                    <strong>Modo Inicial Identificado:</strong> __MODO_INICIAL__<br>
                    <strong>Acción Requerida:</strong> __ACCION_REQUERIDA__<br>
                    <strong>Evaluador:</strong> __EVALUADOR_FULL__
                </td>
            </tr>
        </table>
    </div>

    <div class="section">
        <div class="section-header">3. FUENTES DE ENERGÍA(S) PARA LA(S) TAREA(S) / MODO SELECCIONADO</div>
        <table class="grid-layout">
            <tr>
                <td>
                    <strong>Fuentes de Energía Identificadas:</strong><br>
                    __ENERGIA__
                </td>
            </tr>
        </table>
    </div>

    <div class="section">
        <div class="section-header">4. Identificación de Peligros y Evaluación de Riesgos (W. Fine)</div>
        <table class="matrix-table">
            <thead>
                <tr>
                    <th rowspan="2">Peligro Identificado</th>
                    <th colspan="4">Evaluación de Riesgos (Inherente)</th>
                    <th rowspan="2">Medidas de control</th>
                    <th colspan="4">Evaluación de Riesgos (Residual)</th>
                </tr>
                <tr>
                    <th>P</th><th>E</th><th>C</th><th>GP</th>
                    <th>P</th><th>E</th><th>C</th><th>GP</th>
                </tr>
            </thead>
            <tbody>
                __FILAS_MATRIZ__
            </tbody>
        </table>
        <div style="font-size: 7pt; padding: 5px; background: #f8fafc; border-top: 1px solid #cbd5e1;">
            <strong>Leyenda:</strong> P: Probabilidad | E: Exposición | C: Consecuencia | GP: Grado de Peligrosidad (P x E x C)
        </div>
    </div>

    __SECCION_VALIDACION__

    <!-- Sección Árbol Eliminada -->

    <div class="section">
        <div class="section-header">__NUM_CONCL__. Conclusiones / Observaciones</div>
        <table class="grid-layout">
            <tr>
                <td>
                    <strong>Observaciones:</strong> __CONCLUSIONES__<br>
                    <strong>Modo Final Validado:</strong> __MODO_FINAL__
                </td>
            </tr>
        </table>
    </div>

    <div class="banner-final-pdf" style="background-color: __COLOR_FINAL__; border: none;">
        <div style="font-size: 8pt; font-weight: bold; color: white; text-transform: uppercase;">Modo Final de INTERVENCIÓN VALIDADO</div>
        <div style="font-size: 14pt; font-weight: 900; color: white;">__MODO_FINAL__</div>
        <div style="font-size: 9pt; font-weight: bold; color: white; margin-top: 5px;">__MENSAJE_SEGURIDAD__</div>
    </div>

    <div style="margin-top: 40px; page-break-inside: avoid;">
        <table style="width: 100%; border-collapse: collapse;">
            <tr>
                <td style="width: 50%; border: 1px solid #cbd5e1; padding: 15px; vertical-align: top;">
                    <strong>Evaluó:</strong><br><br>
                    __EVALUADOR_FULL__
                </td>
                <td style="width: 50%; border: 1px solid #cbd5e1; padding: 15px; vertical-align: top;">
                    <strong>Aprobó:</strong><br><br>
                    <div style="border-bottom: 1px solid black; width: 80%; margin-top: 20px;"></div>
                    <div style="font-size: 8pt; margin-top: 5px;">Firma y Aclaración</div>
                </td>
            </tr>
        </table>
    </div>
</body>
</html>
"""

async def generate_pdf(html_content, output_path):
    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            args=[
                "--no-sandbox",
                "--disable-dev-shm-usage",
                "--font-render-hinting=none"
            ]
        )
        page = await browser.new_page(viewport={"width": 1240, "height": 1754})
        await page.emulate_media(media="screen")
        await page.set_content(html_content, wait_until="load")
        await page.wait_for_load_state("networkidle")
        await page.evaluate("() => document.fonts.ready.then(() => true)")
        await page.pdf(
            path=output_path,
            format="A4",
            print_background=True,
            prefer_css_page_size=True,
            margin={"top": "10mm", "right": "8mm", "bottom": "10mm", "left": "8mm"}
        )
        await browser.close()

# --- COMPILACIÓN ---
st.markdown("---")
if st.button("✅ COMPILAR REPORTE (PDF, Word y Excel)"):
    if not lista_peligros_final:
        st.error("Por favor, selecciona al menos un peligro para compilar el informe.")
    else:
        # Mapeo de variables para el reporte
        area_proceso = sitio
        linea_seleccionada = linea
        lista_tareas_final = ", ".join(tareas) if isinstance(tareas, list) else tareas
        
        with st.spinner("Procesando reporte y matriz..."):
            logo_b64 = file_to_base64("arca.png")
            logo_tag = f'<img src="data:image/png;base64,{logo_b64}" style="max-height:45px;">' if logo_b64 else '<h2>ARCA</h2>'
            
            equip_b64 = get_base64_image(equip_file) if equip_file else st.session_state.get("equip_photo_b64")
            equip_tag = f'<img src="data:image/png;base64,{equip_b64}">' if equip_b64 else '<p style="font-size:7pt; color:#64748b;">[Foto no cargada en sistema]</p>'
            
            arbol_b64 = file_to_base64("arbol.png")
            arbol_tag = f'<img src="data:image/png;base64,{arbol_b64}">' if arbol_b64 else '<p style="font-size:7pt; color:#ef4444;">[Falta archivo fijo arbol.png]</p>'
            
            # Construcción estricta de las filas con todo el desglose analítico de Fine
            html_rows = ""
            for item in evaluaciones_fine:
                _, col_r = clasificar_fine(item["gp_r"])
                _, col_i = clasificar_fine(item["gp_i"])
                
                html_rows += f"""
                <tr>
                    <td class="left-align" style="font-weight:bold; width:18%;">{item['peligro']}</td>
                    <td>{item['p_i']}</td>
                    <td>{item['e_i']}</td>
                    <td>{item['c_i']}</td>
                    <td style="background:#f8fafc; font-weight:bold; color:{col_i};">{item['gp_i']}</td>
                    <td class="left-align" style="color:#B51E2D; font-size:7.5pt; width:25%;">{item['controles'] if item['controles'] else 'Sin mitigantes'}</td>
                    <td>{item['p_r']}</td>
                    <td>{item['e_r']}</td>
                    <td>{item['c_r']}</td>
                    <td style="background:#f8fafc; font-weight:bold; color:{col_r};">{item['gp_r']}</td>
                </tr>
                """
            
            html_f = HTML_TEMPLATE
            html_f = html_f.replace("__LOGO_HTML__", logo_tag).replace("__FOTO_EQUIPO__", equip_tag).replace("__ARBOL_LOTO__", arbol_tag)
            html_f = html_f.replace("__PLANTA__", sitio).replace("__FECHA_ACTUAL__", datetime.datetime.now().strftime("%d/%m/%Y"))
            html_f = html_f.replace("__NEGOCIO__", negocio).replace("__TIPO_SITIO__", tipo_sitio)
            html_f = html_f.replace("__AREA__", area_sector).replace("__LINEA__", linea)
            
            equipo_desc = maquina if area_sector == "Producción" else subsector_equipo
            html_f = html_f.replace("__EQUIPO_DESC__", equipo_desc)
            
            html_f = html_f.replace("__FABRICANTE__", fabricante).replace("__MODELO__", modelo).replace("__ANIO__", str(anio))
            html_f = html_f.replace("__CLASIFICACION__", clasificacion)
            html_f = html_f.replace("__FRECUENCIA__", frecuencia).replace("__DURACION__", duracion)
            
            # Lógica de numeración y validación en PDF
            if modo_inicial in ["Modo 1", "Modo 2"]:
                res_val = ""
                if modo_final == "Modo 3" and modo_inicial != "Modo 3":
                    razon = "el riesgo residual es no aceptable (GP > 20)" if riesgo_no_aceptable else "no se garantiza un control fiable"
                    res_val = f"Como {razon}, entonces el modo convalidado es Modo 3."
                else:
                    res_val = "Como el riesgo residual es aceptable (GP < 20) y se garantiza un control fiable, entonces se mantiene el modo Inicial seleccionado."
                
                val_html = f"""
                <div class="section">
                    <div class="section-header">5. Validación del Modo Inicial</div>
                    <table class="grid-layout">
                        <tr>
                            <td>
                                <strong>{pregunta if 'pregunta' in locals() else '¿Se garantiza control fiable para prevenir rearranque?'}</strong><br>
                                <strong>Respuesta:</strong> {control_fiable if 'control_fiable' in locals() else 'N/A'}<br>
                                <strong>Resultado:</strong> {res_val}
                            </td>
                        </tr>
                    </table>
                </div>
                """
                html_f = html_f.replace("__SECCION_VALIDACION__", val_html)
                html_f = html_f.replace("__NUM_CONCL__", "6")
            else:
                html_f = html_f.replace("__SECCION_VALIDACION__", "")
                html_f = html_f.replace("__NUM_CONCL__", "5")
            
            # Sección 2: Identificación de Modo Inicial
            html_f = html_f.replace("__NIVEL_ACCESO__", nivel_acceso)
            html_f = html_f.replace("__MODO_INICIAL__", modo_inicial)
            
            if modo_inicial in ["Modo 1", "Modo 2"]:
                accion_pdf = "*Modo sujeto a validación en 5. VALIDACIÓN DEL MODO INICIAL"
            else:
                accion_pdf = mensajes_seguridad.get(modo_inicial, "N/A")
            html_f = html_f.replace("__ACCION_REQUERIDA__", accion_pdf)
            
            # Formatear tareas para el PDF (separadas por coma)
            tareas_txt = ", ".join(lista_tareas_combinadas) if lista_tareas_combinadas else "No parametrizadas"
            html_f = html_f.replace("__TAREAS__", tareas_txt)
            
            # Formatear energías para el PDF
            energia_txt = ", ".join([f"{k} ({v})" for k, v in energias_seleccionadas.items()]) if energias_seleccionadas else "Ninguna identificada"
            html_f = html_f.replace("__ENERGIA__", energia_txt)
            
            evaluador_full = f"{evaluador} - {puesto_evaluador} - {fecha_evaluacion.strftime('%d/%m/%Y')}"
            html_f = html_f.replace("__EVALUADOR_FULL__", evaluador_full).replace("__FILAS_MATRIZ__", html_rows)
            html_f = html_f.replace("__CONCLUSIONES__", medidas_conclusiones).replace("__MODO_FINAL__", modo_final.upper())
            html_f = html_f.replace("__COLOR_FINAL__", colores_modos.get(modo_final, "#B51E2D"))
            
            # Añadir mensaje de seguridad dinámico al PDF
            msg_pdf = mensajes_seguridad.get(modo_final, "")
            html_f = html_f.replace("__MENSAJE_SEGURIDAD__", msg_pdf)
            
            pdf_out = f"Ficha_Modos_Integral_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            asyncio.run(generate_pdf(html_f, pdf_out))
            
            # --- GENERACIÓN DE EXCEL TÉCNICO AVANZADO, AJUSTADO AL PDF ---
            from openpyxl.drawing.image import Image as XLImage
            from openpyxl.utils.cell import range_boundaries
            from openpyxl import Workbook
            from PIL import Image as PILImage
            import math

            wb = Workbook()
            ws = wb.active
            ws.title = "Ficha Modos Integral"

            # Mantiene vivos los buffers de imágenes hasta guardar el archivo.
            _excel_image_buffers = []

            # ------------------------------------------------------------------
            # Configuración visual y de impresión: A4, similar al PDF generado.
            # ------------------------------------------------------------------
            ws.sheet_view.showGridLines = False
            ws.freeze_panes = "A6"
            ws.page_setup.orientation = "portrait"
            ws.page_setup.paperSize = ws.PAPERSIZE_A4
            ws.page_setup.fitToWidth = 1
            ws.page_setup.fitToHeight = 0
            ws.sheet_properties.pageSetUpPr.fitToPage = True
            ws.page_margins.left = 0.25
            ws.page_margins.right = 0.25
            ws.page_margins.top = 0.35
            ws.page_margins.bottom = 0.35
            ws.page_margins.header = 0.10
            ws.page_margins.footer = 0.10
            ws.print_title_rows = "1:4"

            # Columnas pensadas para replicar el PDF:
            # A:E = texto / F:J = imagen o segunda mitad; matriz Fine en A:J.
            column_widths = {
                "A": 23, "B": 8, "C": 8, "D": 8, "E": 8,
                "F": 32, "G": 8, "H": 8, "I": 8, "J": 8,
            }
            for col, width in column_widths.items():
                ws.column_dimensions[col].width = width

            # Paleta PDF.
            black_fill = PatternFill(start_color="000000", end_color="000000", fill_type="solid")
            dark_fill = PatternFill(start_color="444444", end_color="444444", fill_type="solid")
            red_fill = PatternFill(start_color="B51E2D", end_color="B51E2D", fill_type="solid")
            gray_fill = PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid")
            light_fill = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
            white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
            white_font = Font(name="Bahnschrift", color="FFFFFF", bold=True, size=9)
            white_font_big = Font(name="Bahnschrift", color="FFFFFF", bold=True, size=14)
            white_font_mid = Font(name="Bahnschrift", color="FFFFFF", bold=True, size=11)
            body_font = Font(name="Bahnschrift", color="0F172A", size=8.5)
            bold_font = Font(name="Bahnschrift", color="0F172A", bold=True, size=8.5)
            small_bold_font = Font(name="Bahnschrift", color="0F172A", bold=True, size=8)
            thin_gray = Side(style="thin", color="CBD5E1")
            thin_black = Side(style="thin", color="000000")
            section_side = Side(style="thin", color="B51E2D")
            grid_border = Border(left=thin_gray, right=thin_gray, top=thin_gray, bottom=thin_gray)
            black_border = Border(left=thin_black, right=thin_black, top=thin_black, bottom=thin_black)
            section_border = Border(left=section_side, right=section_side, top=section_side, bottom=section_side)

            def _hex_color(value, fallback="B51E2D"):
                """Normaliza colores tipo #B51E2D para PatternFill/Font."""
                if not value:
                    return fallback
                value = str(value).replace("#", "").strip().upper()
                return value if len(value) == 6 else fallback

            def _style_range(range_ref, fill=None, font=None, border=None, alignment=None):
                """Aplica estilo a todas las celdas de un rango, incluso si se va a combinar."""
                min_col, min_row, max_col, max_row = range_boundaries(range_ref)
                for row_cells in ws.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
                    for cell in row_cells:
                        if fill is not None:
                            cell.fill = fill
                        if font is not None:
                            cell.font = font
                        if border is not None:
                            cell.border = border
                        if alignment is not None:
                            cell.alignment = alignment

            def _merge(range_ref, value="", fill=None, font=None, border=None, alignment=None):
                _style_range(range_ref, fill=fill, font=font, border=border, alignment=alignment)
                ws.merge_cells(range_ref)
                first_cell = range_ref.split(":")[0]
                ws[first_cell] = value
                if font is not None:
                    ws[first_cell].font = font
                if alignment is not None:
                    ws[first_cell].alignment = alignment
                if fill is not None:
                    ws[first_cell].fill = fill
                if border is not None:
                    ws[first_cell].border = border
                return ws[first_cell]

            def _text_height(text, chars_per_line=120, line_height=15, minimum=24, maximum=130):
                text = "" if text is None else str(text)
                manual_lines = text.count("\n") + 1
                estimated_lines = max(manual_lines, math.ceil(len(text) / max(1, chars_per_line)))
                return max(minimum, min(maximum, estimated_lines * line_height + 6))

            def _section_header(title):
                row = ws.max_row + 1
                if row > 5:
                    ws.row_dimensions[row].height = 5
                    row += 1
                _merge(
                    f"A{row}:J{row}",
                    title.upper(),
                    fill=red_fill,
                    font=white_font,
                    border=section_border,
                    alignment=Alignment(horizontal="left", vertical="center"),
                )
                ws.row_dimensions[row].height = 20
                return row

            def _full_width_box(text, chars_per_line=135, fill=white_fill, font=body_font):
                row = ws.max_row + 1
                _merge(
                    f"A{row}:J{row}",
                    text,
                    fill=fill,
                    font=font,
                    border=grid_border,
                    alignment=Alignment(horizontal="left", vertical="top", wrap_text=True),
                )
                ws.row_dimensions[row].height = _text_height(text, chars_per_line=chars_per_line)
                return row

            def _insert_image_fit(raw_bytes, anchor_cell, max_width_px, max_height_px):
                """Inserta una imagen ajustada proporcionalmente dentro de una caja definida."""
                try:
                    pil_img = PILImage.open(io.BytesIO(raw_bytes))
                    pil_img.thumbnail((max_width_px, max_height_px), PILImage.LANCZOS)
                    if pil_img.mode not in ("RGB", "RGBA"):
                        pil_img = pil_img.convert("RGB")
                    out = io.BytesIO()
                    pil_img.save(out, format="PNG")
                    out.seek(0)
                    _excel_image_buffers.append(out)
                    xl_img = XLImage(out)
                    xl_img.width = pil_img.width
                    xl_img.height = pil_img.height
                    ws.add_image(xl_img, anchor_cell)
                    return True
                except Exception:
                    return False

            # ------------------------------------------------------------------
            # 1. Encabezado ISO/SOP, replicando la tabla superior del PDF.
            # ------------------------------------------------------------------
            for r in range(1, 5):
                ws.row_dimensions[r].height = 18

            _merge("A1:B4", "", fill=white_fill, font=body_font, border=black_border, alignment=Alignment(horizontal="center", vertical="center"))
            if os.path.exists("arca.png"):
                try:
                    with open("arca.png", "rb") as logo_file:
                        _insert_image_fit(logo_file.read(), "A1", 125, 58)
                except Exception:
                    ws["A1"] = "ARCA"
                    ws["A1"].font = bold_font
            else:
                ws["A1"] = "ARCA"
                ws["A1"].font = bold_font

            _merge("C1:H2", "CONTROL DE ENERGÍAS PELIGROSAS", fill=black_fill, font=white_font_big, border=black_border, alignment=Alignment(horizontal="center", vertical="center", wrap_text=True))
            _merge("C3:H3", f"{negocio} - {sitio}", fill=black_fill, font=white_font, border=black_border, alignment=Alignment(horizontal="center", vertical="center", wrap_text=True))
            _merge("C4:H4", "MODOS DE INTERVENCIÓN Y EVALUACIÓN DE RIESGOS", fill=dark_fill, font=white_font_mid, border=black_border, alignment=Alignment(horizontal="center", vertical="center", wrap_text=True))
            _merge("I1:J1", "CÓDIGO:", fill=gray_fill, font=small_bold_font, border=black_border, alignment=Alignment(horizontal="left", vertical="center"))
            _merge("I2:J2", "REVISIÓN:", fill=gray_fill, font=small_bold_font, border=black_border, alignment=Alignment(horizontal="left", vertical="center"))
            _merge("I3:J4", f"FECHA: {datetime.datetime.now().strftime('%d/%m/%Y')}", fill=gray_fill, font=small_bold_font, border=black_border, alignment=Alignment(horizontal="left", vertical="center"))

            # Fila técnica de separación, igual al margen visual del PDF.
            ws.row_dimensions[5].height = 8

            # ------------------------------------------------------------------
            # 2. Identificación: texto a la izquierda, foto encapsulada a la derecha.
            # ------------------------------------------------------------------
            _section_header("1. Identificación de Equipo y Tareas")
            start = ws.max_row + 1
            end = start + 7
            for r in range(start, end + 1):
                ws.row_dimensions[r].height = 21

            info_equipo = (
                f"Negocio: {negocio} | Sitio: {sitio}\n"
                f"Ubicación: {area_sector} — {linea}\n"
                f"Equipo/Máquina: {equipo_desc}\n"
                f"Fabricante: {fabricante} | Modelo: {modelo} | Año: {anio}\n"
                f"Clasificación de tarea: {clasificacion}\n"
                f"Frecuencia: {frecuencia} | Duración: {duracion}\n"
                f"Tarea(s):\n{tareas_txt}"
            )
            _merge(f"A{start}:E{end}", info_equipo, fill=white_fill, font=body_font, border=grid_border, alignment=Alignment(horizontal="left", vertical="top", wrap_text=True))
            _merge(f"F{start}:J{start}", "REGISTRO VISUAL DEL EQUIPO:", fill=white_fill, font=bold_font, border=grid_border, alignment=Alignment(horizontal="center", vertical="center"))
            _merge(f"F{start + 1}:J{end}", "", fill=white_fill, font=body_font, border=grid_border, alignment=Alignment(horizontal="center", vertical="center"))

            raw_equipment_photo = None
            if equip_file:
                try:
                    raw_equipment_photo = equip_file.getvalue()
                except Exception:
                    raw_equipment_photo = None
            elif st.session_state.get("equip_photo_b64"):
                try:
                    raw_equipment_photo = base64.b64decode(st.session_state["equip_photo_b64"])
                except Exception:
                    raw_equipment_photo = None

            if raw_equipment_photo:
                inserted = _insert_image_fit(raw_equipment_photo, f"F{start + 1}", 390, 138)
                if not inserted:
                    ws[f"F{start + 1}"] = "[Foto no cargada correctamente]"
            else:
                ws[f"F{start + 1}"] = "[Foto no cargada en sistema]"
                ws[f"F{start + 1}"].font = Font(name="Bahnschrift", color="64748B", italic=True, size=8)

            # ------------------------------------------------------------------
            # 3. Modo inicial.
            # ------------------------------------------------------------------
            _section_header("2. Identificación de Modo Inicial de Intervención")
            modo_text = (
                f"Nivel de Acceso / Exposición Requerido: {nivel_acceso}\n"
                f"Modo Inicial Identificado: {modo_inicial}\n"
                f"Acción Requerida: {accion_pdf}\n"
                f"Evaluador: {evaluador_full}"
            )
            _full_width_box(modo_text, chars_per_line=130)

            # ------------------------------------------------------------------
            # 4. Fuentes de energía.
            # ------------------------------------------------------------------
            _section_header("3. Fuentes de Energía(s) para la(s) Tarea(s) / Modo Seleccionado")
            energia_text = f"Fuentes de Energía Identificadas:\n{energia_txt}"
            _full_width_box(energia_text, chars_per_line=130)

            # ------------------------------------------------------------------
            # 5. Matriz William Fine, con encabezado agrupado como en el PDF.
            # ------------------------------------------------------------------
            _section_header("4. Identificación de Peligros y Evaluación de Riesgos (W. Fine)")
            h1 = ws.max_row + 1
            h2 = h1 + 1
            ws.row_dimensions[h1].height = 20
            ws.row_dimensions[h2].height = 18
            _merge(f"A{h1}:A{h2}", "Peligro Identificado", fill=dark_fill, font=white_font, border=black_border, alignment=Alignment(horizontal="center", vertical="center", wrap_text=True))
            _merge(f"B{h1}:E{h1}", "Evaluación de Riesgos (Inherente)", fill=dark_fill, font=white_font, border=black_border, alignment=Alignment(horizontal="center", vertical="center", wrap_text=True))
            _merge(f"F{h1}:F{h2}", "Medidas de control", fill=dark_fill, font=white_font, border=black_border, alignment=Alignment(horizontal="center", vertical="center", wrap_text=True))
            _merge(f"G{h1}:J{h1}", "Evaluación de Riesgos (Residual)", fill=dark_fill, font=white_font, border=black_border, alignment=Alignment(horizontal="center", vertical="center", wrap_text=True))
            for col, label in zip(["B", "C", "D", "E", "G", "H", "I", "J"], ["P", "E", "C", "GP", "P", "E", "C", "GP"]):
                cell = ws[f"{col}{h2}"]
                cell.value = label
                cell.fill = dark_fill
                cell.font = white_font
                cell.border = black_border
                cell.alignment = Alignment(horizontal="center", vertical="center")

            for item in evaluaciones_fine:
                row = ws.max_row + 1
                values = [
                    item.get("peligro", ""), item.get("p_i", ""), item.get("e_i", ""), item.get("c_i", ""), item.get("gp_i", ""),
                    item.get("controles", "") if item.get("controles", "") else "Sin mitigantes",
                    item.get("p_r", ""), item.get("e_r", ""), item.get("c_r", ""), item.get("gp_r", ""),
                ]
                for col_index, value in enumerate(values, start=1):
                    cell = ws.cell(row=row, column=col_index, value=value)
                    cell.font = body_font
                    cell.border = grid_border
                    cell.alignment = Alignment(horizontal="center", vertical="top", wrap_text=True)
                    if col_index in [1, 6]:
                        cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                    if col_index in [5, 10]:
                        try:
                            _, risk_color = clasificar_fine(float(value))
                            cell.font = Font(name="Bahnschrift", color=_hex_color(risk_color), bold=True, size=8.5)
                        except Exception:
                            cell.font = Font(name="Bahnschrift", color="0F172A", bold=True, size=8.5)
                        cell.fill = light_fill
                peligro_len = len(str(values[0]))
                controles_len = len(str(values[5]))
                estimated_lines = max(math.ceil(peligro_len / 28), math.ceil(controles_len / 42), 1)
                ws.row_dimensions[row].height = min(100, max(28, estimated_lines * 16 + 8))

            legend_row = ws.max_row + 1
            _merge(
                f"A{legend_row}:J{legend_row}",
                "Leyenda: P: Probabilidad | E: Exposición | C: Consecuencia | GP: Grado de Peligrosidad (P x E x C)",
                fill=light_fill,
                font=Font(name="Bahnschrift", color="0F172A", bold=True, size=7),
                border=grid_border,
                alignment=Alignment(horizontal="left", vertical="center", wrap_text=True),
            )
            ws.row_dimensions[legend_row].height = 18

            # ------------------------------------------------------------------
            # 6. Validación del modo inicial, cuando corresponde.
            # ------------------------------------------------------------------
            if modo_inicial in ["Modo 1", "Modo 2"]:
                _section_header("5. Validación del Modo Inicial")
                validacion_text = (
                    f"{pregunta if 'pregunta' in locals() else '¿Se garantiza control fiable para prevenir rearranque?'}\n"
                    f"Respuesta: {control_fiable if 'control_fiable' in locals() else 'N/A'}\n"
                    f"Resultado: {res_val}"
                )
                _full_width_box(validacion_text, chars_per_line=130)
                concl_number = "6"
            else:
                concl_number = "5"

            # ------------------------------------------------------------------
            # 7. Conclusiones, banner final y firmas, igual que el PDF.
            # ------------------------------------------------------------------
            _section_header(f"{concl_number}. Conclusiones / Observaciones")
            conclusiones_text = (
                f"Observaciones: {medidas_conclusiones}\n"
                f"Modo Final Validado: {modo_final.upper()}"
            )
            _full_width_box(conclusiones_text, chars_per_line=130)

            banner_color = _hex_color(colores_modos.get(modo_final, "#B51E2D"))
            banner_fill = PatternFill(start_color=banner_color, end_color=banner_color, fill_type="solid")
            banner_row = ws.max_row + 1
            banner_text = f"Modo Final de INTERVENCIÓN VALIDADO\n{modo_final.upper()}\n{msg_pdf}"
            _merge(
                f"A{banner_row}:J{banner_row}",
                banner_text,
                fill=banner_fill,
                font=Font(name="Bahnschrift", color="FFFFFF", bold=True, size=12),
                border=Border(),
                alignment=Alignment(horizontal="center", vertical="center", wrap_text=True),
            )
            ws.row_dimensions[banner_row].height = 62

            ws.row_dimensions[ws.max_row + 1].height = 10
            firma_row = ws.max_row + 2
            _merge(f"A{firma_row}:E{firma_row + 2}", f"Evaluó:\n\n{evaluador_full}", fill=white_fill, font=body_font, border=grid_border, alignment=Alignment(horizontal="left", vertical="top", wrap_text=True))
            _merge(f"F{firma_row}:J{firma_row + 2}", "Aprobó:\n\n______________________________\nFirma y Aclaración", fill=white_fill, font=body_font, border=grid_border, alignment=Alignment(horizontal="left", vertical="top", wrap_text=True))
            for r in range(firma_row, firma_row + 3):
                ws.row_dimensions[r].height = 28

            # Aplicar fuente y alineación segura a todo el rango usado.
            for row_cells in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=10):
                for cell in row_cells:
                    if cell.font is None:
                        cell.font = body_font
                    if cell.alignment is None:
                        cell.alignment = Alignment(vertical="top", wrap_text=True)

            ws.print_area = f"A1:J{ws.max_row}"

            excel_out = f"Ficha_Modos_Integral_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            wb.save(excel_out)

                    # --- GENERACIÓN DE WORD (.DOCX) AJUSTADO AL PDF ---
            # Requiere: pip install python-docx
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from PIL import Image as PILImage
            import math
            import io

            doc = Document()
            section = doc.sections[0]
            section.page_width = Inches(8.27)     # A4 vertical
            section.page_height = Inches(11.69)
            section.left_margin = Inches(0.32)
            section.right_margin = Inches(0.32)
            section.top_margin = Inches(0.40)
            section.bottom_margin = Inches(0.40)
            section.header_distance = Inches(0.10)
            section.footer_distance = Inches(0.10)

            # Paleta equivalente al PDF.
            WORD_BLACK = "000000"
            WORD_DARK = "444444"
            WORD_RED = "B51E2D"
            WORD_GRAY = "F1F5F9"
            WORD_LIGHT = "F8FAFC"
            WORD_GRID = "CBD5E1"
            WORD_TEXT = "0F172A"
            WORD_WHITE = "FFFFFF"

            styles = doc.styles
            styles["Normal"].font.name = "Bahnschrift"
            styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Bahnschrift")
            styles["Normal"].font.size = Pt(8.5)
            styles["Normal"].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

            _docx_image_buffers = []

            def _clean_hex(value, fallback=WORD_RED):
                if not value:
                    return fallback
                value = str(value).replace("#", "").strip().upper()
                return value if len(value) == 6 else fallback

            def _rgb(hex_color):
                hex_color = _clean_hex(hex_color, WORD_TEXT)
                return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

            def _set_cell_shading(cell, fill):
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = tc_pr.find(qn("w:shd"))
                if shd is None:
                    shd = OxmlElement("w:shd")
                    tc_pr.append(shd)
                shd.set(qn("w:fill"), _clean_hex(fill))

            def _set_cell_width(cell, width_twips):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(width_twips))
                tc_w.set(qn("w:type"), "dxa")

            def _set_cell_borders(cell, color=WORD_GRID, size="6"):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_borders = tc_pr.first_child_found_in("w:tcBorders")
                if tc_borders is None:
                    tc_borders = OxmlElement("w:tcBorders")
                    tc_pr.append(tc_borders)
                for edge in ("top", "left", "bottom", "right"):
                    tag = "w:{}".format(edge)
                    element = tc_borders.find(qn(tag))
                    if element is None:
                        element = OxmlElement(tag)
                        tc_borders.append(element)
                    element.set(qn("w:val"), "single")
                    element.set(qn("w:sz"), size)
                    element.set(qn("w:space"), "0")
                    element.set(qn("w:color"), _clean_hex(color, WORD_GRID))

            def _set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
                tc = cell._tc
                tc_pr = tc.get_or_add_tcPr()
                tc_mar = tc_pr.first_child_found_in("w:tcMar")
                if tc_mar is None:
                    tc_mar = OxmlElement("w:tcMar")
                    tc_pr.append(tc_mar)
                for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
                    node = tc_mar.find(qn(f"w:{margin_name}"))
                    if node is None:
                        node = OxmlElement(f"w:{margin_name}")
                        tc_mar.append(node)
                    node.set(qn("w:w"), str(value))
                    node.set(qn("w:type"), "dxa")

            def _set_table_width_pct(table, pct=5000):
                tbl_pr = table._tbl.tblPr
                tbl_w = tbl_pr.find(qn("w:tblW"))
                if tbl_w is None:
                    tbl_w = OxmlElement("w:tblW")
                    tbl_pr.append(tbl_w)
                tbl_w.set(qn("w:w"), str(pct))
                tbl_w.set(qn("w:type"), "pct")

            def _set_table_fixed(table):
                tbl_pr = table._tbl.tblPr
                layout = tbl_pr.find(qn("w:tblLayout"))
                if layout is None:
                    layout = OxmlElement("w:tblLayout")
                    tbl_pr.append(layout)
                layout.set(qn("w:type"), "fixed")

            def _format_paragraph(paragraph, align=None, space_after=0, space_before=0, line_spacing=1.0):
                if align is not None:
                    paragraph.alignment = align
                paragraph.paragraph_format.space_after = Pt(space_after)
                paragraph.paragraph_format.space_before = Pt(space_before)
                paragraph.paragraph_format.line_spacing = line_spacing

            def _clear_cell(cell):
                cell.text = ""
                if not cell.paragraphs:
                    cell.add_paragraph()

            def _write_text(cell, text, bold=False, color=WORD_TEXT, size=8.5, align=None, uppercase=False):
                _clear_cell(cell)
                lines = str(text if text is not None else "").split("\n")
                for idx, line in enumerate(lines):
                    paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
                    _format_paragraph(paragraph, align=align, space_after=0, line_spacing=1.0)
                    run = paragraph.add_run(line.upper() if uppercase else line)
                    run.bold = bold
                    run.font.name = "Bahnschrift"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Bahnschrift")
                    run.font.size = Pt(size)
                    run.font.color.rgb = _rgb(color)

            def _add_label_line(cell, label, value="", first=False):
                paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
                _format_paragraph(paragraph, space_after=0, line_spacing=1.0)
                run_label = paragraph.add_run(label)
                run_label.bold = True
                run_label.font.name = "Bahnschrift"
                run_label._element.rPr.rFonts.set(qn("w:eastAsia"), "Bahnschrift")
                run_label.font.size = Pt(8.5)
                run_label.font.color.rgb = _rgb(WORD_TEXT)
                run_value = paragraph.add_run(str(value if value is not None else ""))
                run_value.font.name = "Bahnschrift"
                run_value._element.rPr.rFonts.set(qn("w:eastAsia"), "Bahnschrift")
                run_value.font.size = Pt(8.5)
                run_value.font.color.rgb = _rgb(WORD_TEXT)

            def _style_all_cells(table, border_color=WORD_GRID, fill=None, valign=WD_CELL_VERTICAL_ALIGNMENT.TOP):
                for row in table.rows:
                    for cell in row.cells:
                        cell.vertical_alignment = valign
                        _set_cell_borders(cell, border_color)
                        _set_cell_margins(cell)
                        if fill:
                            _set_cell_shading(cell, fill)

            def _section_header(title):
                table = doc.add_table(rows=1, cols=1)
                table.autofit = False
                _set_table_width_pct(table, 5000)
                _set_table_fixed(table)
                cell = table.cell(0, 0)
                # Ancho útil A4 con márgenes laterales de 0.32 pulgadas: 7.63 in = 10987 twips aprox.
                # Esto evita que Word ajuste la barra roja al contenido y deje blanco a la derecha.
                _set_cell_width(cell, 10980)
                _set_cell_shading(cell, WORD_RED)
                _set_cell_borders(cell, WORD_RED, size="8")
                _set_cell_margins(cell, top=90, start=140, bottom=90, end=140)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _write_text(cell, title, bold=True, color=WORD_WHITE, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, uppercase=True)
                return table

            def _single_box():
                table = doc.add_table(rows=1, cols=1)
                table.autofit = False
                _set_table_width_pct(table)
                _style_all_cells(table, WORD_GRID)
                return table, table.cell(0, 0)

            def _image_buffer_fit(raw_bytes, max_width_px, max_height_px):
                pil_img = PILImage.open(io.BytesIO(raw_bytes))
                pil_img.thumbnail((max_width_px, max_height_px), PILImage.LANCZOS)
                if pil_img.mode not in ("RGB", "RGBA"):
                    pil_img = pil_img.convert("RGB")
                output = io.BytesIO()
                pil_img.save(output, format="PNG")
                output.seek(0)
                _docx_image_buffers.append(output)
                return output, pil_img.width, pil_img.height

            def _add_picture_to_cell(cell, raw_bytes, max_width_in=3.25, max_height_in=1.65):
                _clear_cell(cell)
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
                try:
                    buffer, width_px, height_px = _image_buffer_fit(raw_bytes, int(max_width_in * 96), int(max_height_in * 96))
                    run = paragraph.add_run()
                    run.add_picture(buffer, width=Inches(width_px / 96))
                    return True
                except Exception:
                    _write_text(cell, "[Foto no cargada correctamente]", color="64748B", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                    return False

            def _add_spacing(points=6):
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(points)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

            # ------------------------------------------------------------------
            # Encabezado ISO/SOP.
            # ------------------------------------------------------------------
            header_table = doc.add_table(rows=1, cols=3)
            header_table.autofit = False
            _set_table_width_pct(header_table)
            _set_table_fixed(header_table)
            _style_all_cells(header_table, WORD_BLACK, valign=WD_CELL_VERTICAL_ALIGNMENT.CENTER)
            _set_cell_width(header_table.cell(0, 0), 1900)
            _set_cell_width(header_table.cell(0, 1), 5250)
            _set_cell_width(header_table.cell(0, 2), 2350)

            logo_cell = header_table.cell(0, 0)
            _clear_cell(logo_cell)
            logo_p = logo_cell.paragraphs[0]
            _format_paragraph(logo_p, align=WD_ALIGN_PARAGRAPH.CENTER)
            if os.path.exists("arca.png"):
                try:
                    logo_p.add_run().add_picture("arca.png", width=Inches(1.25))
                except Exception:
                    _write_text(logo_cell, "ARCA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                _write_text(logo_cell, "ARCA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

            title_cell = header_table.cell(0, 1)
            _set_cell_shading(title_cell, WORD_BLACK)
            _clear_cell(title_cell)
            p1 = title_cell.paragraphs[0]
            _format_paragraph(p1, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
            r1 = p1.add_run("CONTROL DE ENERGÍAS PELIGROSAS")
            r1.bold = True
            r1.font.name = "Bahnschrift"
            r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Bahnschrift")
            r1.font.size = Pt(12)
            r1.font.color.rgb = _rgb(WORD_WHITE)
            p2 = title_cell.add_paragraph()
            _format_paragraph(p2, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
            r2 = p2.add_run(f"{negocio} - {sitio}")
            r2.font.name = "Bahnschrift"
            r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Bahnschrift")
            r2.font.size = Pt(10)
            r2.font.color.rgb = _rgb("CCCCCC")
            p3 = title_cell.add_paragraph()
            _format_paragraph(p3, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
            r3 = p3.add_run("MODOS DE INTERVENCIÓN Y EVALUACIÓN DE RIESGOS")
            r3.bold = True
            r3.font.name = "Bahnschrift"
            r3._element.rPr.rFonts.set(qn("w:eastAsia"), "Bahnschrift")
            r3.font.size = Pt(11)
            r3.font.color.rgb = _rgb(WORD_WHITE)

            info_cell = header_table.cell(0, 2)
            _set_cell_shading(info_cell, "EEEEEE")
            _clear_cell(info_cell)
            _add_label_line(info_cell, "CÓDIGO: ", "", first=True)
            _add_label_line(info_cell, "REVISIÓN: ", "")
            _add_label_line(info_cell, "FECHA: ", datetime.datetime.now().strftime("%d/%m/%Y"))
            _add_spacing(5)

            # ------------------------------------------------------------------
            # 1. Identificación de equipo y tareas.
            # ------------------------------------------------------------------
            _section_header("1. Identificación de Equipo y Tareas")
            ident_table = doc.add_table(rows=1, cols=2)
            ident_table.autofit = False
            _set_table_width_pct(ident_table)
            _set_table_fixed(ident_table)
            _style_all_cells(ident_table, WORD_GRID)
            _set_cell_width(ident_table.cell(0, 0), 4750)
            _set_cell_width(ident_table.cell(0, 1), 4750)

            left_cell = ident_table.cell(0, 0)
            _clear_cell(left_cell)
            _add_label_line(left_cell, "Negocio: ", f"{negocio} | Sitio: {sitio}", first=True)
            _add_label_line(left_cell, "Ubicación: ", f"{area_sector} — {linea}")
            _add_label_line(left_cell, "Equipo/Máquina: ", equipo_desc)
            _add_label_line(left_cell, "Fabricante: ", f"{fabricante} | Modelo: {modelo} | Año: {anio}")
            _add_label_line(left_cell, "Clasificación de tarea: ", clasificacion)
            _add_label_line(left_cell, "Frecuencia: ", f"{frecuencia} | Duración: {duracion}")
            _add_label_line(left_cell, "Tarea(s):", "")
            p_tasks = left_cell.add_paragraph(str(tareas_txt))
            _format_paragraph(p_tasks, space_after=0, line_spacing=1.0)
            for run in p_tasks.runs:
                run.font.name = "Bahnschrift"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Bahnschrift")
                run.font.size = Pt(8.5)
                run.font.color.rgb = _rgb(WORD_TEXT)

            right_cell = ident_table.cell(0, 1)
            _clear_cell(right_cell)
            p_caption = right_cell.paragraphs[0]
            _format_paragraph(p_caption, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
            r_caption = p_caption.add_run("REGISTRO VISUAL DEL EQUIPO:")
            r_caption.bold = True
            r_caption.font.name = "Bahnschrift"
            r_caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Bahnschrift")
            r_caption.font.size = Pt(8.5)
            r_caption.font.color.rgb = _rgb(WORD_TEXT)
            raw_equipment_photo_word = None
            if equip_file:
                try:
                    raw_equipment_photo_word = equip_file.getvalue()
                except Exception:
                    raw_equipment_photo_word = None
            elif st.session_state.get("equip_photo_b64"):
                try:
                    raw_equipment_photo_word = base64.b64decode(st.session_state["equip_photo_b64"])
                except Exception:
                    raw_equipment_photo_word = None
            if raw_equipment_photo_word:
                pic_paragraph = right_cell.add_paragraph()
                _format_paragraph(pic_paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
                try:
                    buffer, width_px, height_px = _image_buffer_fit(raw_equipment_photo_word, 310, 155)
                    pic_paragraph.add_run().add_picture(buffer, width=Inches(width_px / 96))
                except Exception:
                    right_cell.add_paragraph("[Foto no cargada correctamente]")
            else:
                p_no = right_cell.add_paragraph("[Foto no cargada en sistema]")
                _format_paragraph(p_no, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

            _add_spacing(5)

            # ------------------------------------------------------------------
            # 2. Identificación de modo inicial.
            # ------------------------------------------------------------------
            _section_header("2. Identificación de Modo Inicial de Intervención")
            mode_table, mode_cell = _single_box()
            _clear_cell(mode_cell)
            _add_label_line(mode_cell, "Nivel de Acceso / Exposición Requerido: ", nivel_acceso, first=True)
            _add_label_line(mode_cell, "Modo Inicial Identificado: ", modo_inicial)
            _add_label_line(mode_cell, "Acción Requerida: ", accion_pdf)
            _add_label_line(mode_cell, "Evaluador: ", evaluador_full)
            _add_spacing(5)

            # ------------------------------------------------------------------
            # 3. Fuentes de energía.
            # ------------------------------------------------------------------
            _section_header("3. Fuentes de Energía(s) para la(s) Tarea(s) / Modo Seleccionado")
            energy_table, energy_cell = _single_box()
            _clear_cell(energy_cell)
            _add_label_line(energy_cell, "Fuentes de Energía Identificadas:", "", first=True)
            p_energy = energy_cell.add_paragraph(str(energia_txt))
            _format_paragraph(p_energy, space_after=0, line_spacing=1.0)
            for run in p_energy.runs:
                run.font.name = "Bahnschrift"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Bahnschrift")
                run.font.size = Pt(8.5)
                run.font.color.rgb = _rgb(WORD_TEXT)
            _add_spacing(5)

            # ------------------------------------------------------------------
            # 4. Matriz William Fine, con encabezado agrupado como el PDF.
            # ------------------------------------------------------------------
            _section_header("4. Identificación de Peligros y Evaluación de Riesgos (W. Fine)")
            matrix_table = doc.add_table(rows=2, cols=10)
            matrix_table.autofit = False
            _set_table_width_pct(matrix_table)
            _set_table_fixed(matrix_table)
            _style_all_cells(matrix_table, WORD_GRID)
            col_twips = [1850, 520, 520, 520, 600, 2500, 520, 520, 520, 600]
            for row in matrix_table.rows:
                for idx, width in enumerate(col_twips):
                    _set_cell_width(row.cells[idx], width)

            h0 = matrix_table.rows[0].cells
            h1 = matrix_table.rows[1].cells
            h0[0].merge(h1[0])
            h0[1].merge(h0[4])
            h0[5].merge(h1[5])
            h0[6].merge(h0[9])
            headers = [
                (matrix_table.cell(0, 0), "Peligro Identificado"),
                (matrix_table.cell(0, 1), "Evaluación de Riesgos (Inherente)"),
                (matrix_table.cell(0, 5), "Medidas de control"),
                (matrix_table.cell(0, 6), "Evaluación de Riesgos (Residual)"),
            ]
            for cell, text in headers:
                _set_cell_shading(cell, WORD_DARK)
                _set_cell_borders(cell, WORD_DARK)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _write_text(cell, text, bold=True, color=WORD_WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            for idx, label in zip([1, 2, 3, 4, 6, 7, 8, 9], ["P", "E", "C", "GP", "P", "E", "C", "GP"]):
                cell = matrix_table.cell(1, idx)
                _set_cell_shading(cell, WORD_DARK)
                _set_cell_borders(cell, WORD_DARK)
                _write_text(cell, label, bold=True, color=WORD_WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

            for item in evaluaciones_fine:
                cells = matrix_table.add_row().cells
                values = [
                    item.get("peligro", ""), item.get("p_i", ""), item.get("e_i", ""), item.get("c_i", ""), item.get("gp_i", ""),
                    item.get("controles", "") if item.get("controles", "") else "Sin mitigantes",
                    item.get("p_r", ""), item.get("e_r", ""), item.get("c_r", ""), item.get("gp_r", ""),
                ]
                for idx, value in enumerate(values):
                    cell = cells[idx]
                    _set_cell_borders(cell, WORD_GRID)
                    _set_cell_margins(cell, top=80, start=80, bottom=80, end=80)
                    align = WD_ALIGN_PARAGRAPH.LEFT if idx in [0, 5] else WD_ALIGN_PARAGRAPH.CENTER
                    color = WORD_TEXT
                    bold = False
                    if idx in [4, 9]:
                        try:
                            _, risk_color = clasificar_fine(float(value))
                            color = _clean_hex(risk_color, WORD_TEXT)
                        except Exception:
                            color = WORD_TEXT
                        bold = True
                        _set_cell_shading(cell, WORD_LIGHT)
                    _write_text(cell, value, bold=bold, color=color, size=7.5, align=align)

            legend_table, legend_cell = _single_box()
            _set_cell_shading(legend_cell, WORD_LIGHT)
            _write_text(
                legend_cell,
                "Leyenda: P: Probabilidad | E: Exposición | C: Consecuencia | GP: Grado de Peligrosidad (P x E x C)",
                bold=True,
                color=WORD_TEXT,
                size=7,
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )
            _add_spacing(5)

            # ------------------------------------------------------------------
            # 5. Validación del modo inicial, cuando corresponde.
            # ------------------------------------------------------------------
            if modo_inicial in ["Modo 1", "Modo 2"]:
                _section_header("5. Validación del Modo Inicial")
                val_table, val_cell = _single_box()
                _clear_cell(val_cell)
                _add_label_line(val_cell, "", pregunta if 'pregunta' in locals() else "¿Se garantiza control fiable para prevenir rearranque?", first=True)
                _add_label_line(val_cell, "Respuesta: ", control_fiable if 'control_fiable' in locals() else "N/A")
                _add_label_line(val_cell, "Resultado: ", res_val)
                concl_number_word = "6"
                _add_spacing(5)
            else:
                concl_number_word = "5"

            # ------------------------------------------------------------------
            # 6. Conclusión, banner final y firmas.
            # ------------------------------------------------------------------
            _section_header(f"{concl_number_word}. Conclusiones / Observaciones")
            concl_table, concl_cell = _single_box()
            _clear_cell(concl_cell)
            _add_label_line(concl_cell, "Observaciones: ", medidas_conclusiones, first=True)
            _add_label_line(concl_cell, "Modo Final Validado: ", modo_final.upper())
            _add_spacing(5)

            banner_color_word = _clean_hex(colores_modos.get(modo_final, "#B51E2D"), WORD_RED)
            banner_table = doc.add_table(rows=1, cols=1)
            banner_table.autofit = False
            _set_table_width_pct(banner_table)
            banner_cell = banner_table.cell(0, 0)
            _set_cell_shading(banner_cell, banner_color_word)
            _set_cell_borders(banner_cell, banner_color_word)
            _set_cell_margins(banner_cell, top=150, start=150, bottom=150, end=150)
            banner_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _clear_cell(banner_cell)
            for idx, (text, size) in enumerate([
                ("Modo Final de INTERVENCIÓN VALIDADO", 8),
                (modo_final.upper(), 14),
                (msg_pdf, 9),
            ]):
                paragraph = banner_cell.paragraphs[0] if idx == 0 else banner_cell.add_paragraph()
                _format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
                run = paragraph.add_run(str(text))
                run.bold = True
                run.font.name = "Bahnschrift"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Bahnschrift")
                run.font.size = Pt(size)
                run.font.color.rgb = _rgb(WORD_WHITE)
            _add_spacing(14)

            firma_table = doc.add_table(rows=1, cols=2)
            firma_table.autofit = False
            _set_table_width_pct(firma_table)
            _set_table_fixed(firma_table)
            _style_all_cells(firma_table, WORD_GRID)
            _set_cell_width(firma_table.cell(0, 0), 4750)
            _set_cell_width(firma_table.cell(0, 1), 4750)
            firma_left = firma_table.cell(0, 0)
            firma_right = firma_table.cell(0, 1)
            _write_text(firma_left, f"Evaluó:\n\n{evaluador_full}", color=WORD_TEXT, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)
            _write_text(firma_right, "Aprobó:\n\n______________________________\nFirma y Aclaración", color=WORD_TEXT, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)

            word_out = f"Ficha_LOTO_Integral_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(word_out)

            # --- GENERAR BORRADOR JSON AL COMPILAR ---
            draft_payload = _build_draft_payload(
                evaluaciones_fine=evaluaciones_fine,
                energias_seleccionadas=energias_seleccionadas,
                lista_peligros_final=lista_peligros_final,
                modo_inicial=modo_inicial,
                modo_final=modo_final,
                uploaded_file=equip_file,
            )

            st.session_state["draft_json_download"] = json.dumps(
                draft_payload,
                ensure_ascii=False,
                indent=4,
            ).encode("utf-8")

            st.session_state["draft_json_filename"] = (
                f"Borrador_LOTO_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            )





            
            with open(pdf_out, "rb") as f:
                st.session_state["pdf_v5"] = f.read()
            with open(excel_out, "rb") as f:
                st.session_state["excel_v5"] = f.read()
            with open(word_out, "rb") as f:
                st.session_state["word_v5"] = f.read()
                
            st.success("🎉 Reporte (PDF + Word + Excel) compilado.")

if "pdf_v5" in st.session_state:
    st.download_button("📥 Descargar Reporte en PDF", st.session_state["pdf_v5"], file_name=f"Ficha_Modos_Integral.pdf", mime="application/pdf")
if "word_v5" in st.session_state:
    st.download_button("📥 Descargar Reporte en Word", st.session_state["word_v5"], file_name=f"Ficha_Modos_Integral.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
if "excel_v5" in st.session_state:
    st.download_button("📥 Descargar Reporte en Excel", st.session_state["excel_v5"], file_name=f"Ficha_Modos_Integral.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
if "draft_json_download" in st.session_state:
    st.download_button("📥 Descargar Borrador para Procedimiento de Modos", st.session_state["draft_json_download"], file_name=f"Borrador_LOTO.json", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheetapplication/json")